"""Convert a Markdown file to a SIL-branded .docx using the bundled template.

Usage:
    python convert.py INPUT.md OUTPUT.docx \
        [--title "Document Title"] \
        [--subtitle "Optional subtitle line"] \
        [--template path/to/template.docx]

If --template is omitted, the sibling `sil-template.docx` is used.
Supports: # / ## / ### / #### headings, numbered lists, bullets (incl. nested),
markdown tables, bold/italic/inline-code, and a final italic-only line.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

BODY_SPACE_AFTER_PT = 8


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _w(tag: str) -> str:
    return f"{{{W}}}{tag}"


class NumberingRegistry:
    """Manages numbering.xml: one abstractNum per list kind, a fresh numId per list instance.

    Each call to new_list(kind) returns a numId that will start its counter at 1,
    so every distinct numbered list in the document auto-numbers independently.
    """

    def __init__(self, doc):
        self.doc = doc
        self.numbering = doc.part.numbering_part.element
        self._used_abs = {
            int(e.get(_w("abstractNumId")))
            for e in self.numbering.findall(_w("abstractNum"))
        }
        self._used_num = {
            int(e.get(_w("numId")))
            for e in self.numbering.findall(_w("num"))
        }
        self._abstract_for_kind: dict[str, int] = {}

    def _next(self, used: set[int]) -> int:
        i = 1
        while i in used:
            i += 1
        used.add(i)
        return i

    def _ensure_abstract(self, kind: str) -> int:
        if kind in self._abstract_for_kind:
            return self._abstract_for_kind[kind]
        # kind refers to the TOP-LEVEL list style. Nested levels always use
        # bullets so a markdown numbered list with nested "-" items renders as
        # "1. … • … ▪ …" without the nested items resetting the outer counter.
        if kind == "decimal":
            fmts = ["decimal", "bullet", "bullet"]
            texts = ["%1.", "\u25e6", "\u25aa"]
        else:
            fmts = ["bullet", "bullet", "bullet"]
            texts = ["\u2022", "\u25e6", "\u25aa"]
        abs_id = self._next(self._used_abs)
        abs_el = OxmlElement("w:abstractNum")
        abs_el.set(_w("abstractNumId"), str(abs_id))
        for lvl_i, (fmt, txt) in enumerate(zip(fmts, texts)):
            lvl = OxmlElement("w:lvl")
            lvl.set(_w("ilvl"), str(lvl_i))
            start = OxmlElement("w:start"); start.set(_w("val"), "1"); lvl.append(start)
            nfmt = OxmlElement("w:numFmt"); nfmt.set(_w("val"), fmt); lvl.append(nfmt)
            lt = OxmlElement("w:lvlText"); lt.set(_w("val"), txt); lvl.append(lt)
            jc = OxmlElement("w:lvlJc"); jc.set(_w("val"), "left"); lvl.append(jc)
            pPr = OxmlElement("w:pPr")
            ind = OxmlElement("w:ind")
            ind.set(_w("left"), str(720 + lvl_i * 360))
            ind.set(_w("hanging"), "360")
            pPr.append(ind); lvl.append(pPr)
            abs_el.append(lvl)
        first_num = self.numbering.find(_w("num"))
        if first_num is not None:
            first_num.addprevious(abs_el)
        else:
            self.numbering.append(abs_el)
        self._abstract_for_kind[kind] = abs_id
        return abs_id

    def new_list(self, kind: str) -> int:
        """Allocate a fresh numId bound to the abstractNum for this kind.

        Each call returns a new numId, so counters restart at 1.
        """
        abs_id = self._ensure_abstract(kind)
        num_id = self._next(self._used_num)
        num_el = OxmlElement("w:num")
        num_el.set(_w("numId"), str(num_id))
        ref = OxmlElement("w:abstractNumId"); ref.set(_w("val"), str(abs_id))
        num_el.append(ref)
        # Force each level's counter to start at 1 for this num instance, so
        # separate lists don't continue counting from a previous list that
        # shared the same abstractNum.
        for lvl in range(3):
            ov = OxmlElement("w:lvlOverride")
            ov.set(_w("ilvl"), str(lvl))
            start_ov = OxmlElement("w:startOverride")
            start_ov.set(_w("val"), "1")
            ov.append(start_ov)
            num_el.append(ov)
        self.numbering.append(num_el)
        return num_id


def apply_list_item(paragraph, num_id: int, level: int) -> None:
    """Attach numPr (numbering reference) to a paragraph so Word auto-numbers it."""
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(_w("val"), str(level)); numPr.append(ilvl)
    nid = OxmlElement("w:numId"); nid.set(_w("val"), str(num_id)); numPr.append(nid)
    pPr.append(numPr)

INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def add_inline(paragraph, text: str) -> None:
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1]); r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1]); r.font.name = "Consolas"
        else:
            paragraph.add_run(part)


def has_style(doc, name: str) -> bool:
    for s in doc.styles:
        if (s.name or "").lower() == name.lower():
            return True
    return False


def style_name(doc, preferred: str, fallback: str = "Normal") -> str:
    return preferred if has_style(doc, preferred) else fallback


def make_heading3_bold(doc) -> None:
    """Add bold to the Heading 3 style definition, plus a bit of space above."""
    if has_style(doc, "Heading 3"):
        s = doc.styles["Heading 3"]
        s.font.bold = True
        s.paragraph_format.space_before = Pt(14)
        s.paragraph_format.space_after = Pt(6)


def set_title_block(doc, title: str | None, subtitle: str | None) -> None:
    """Replace Title / Subtitle paragraph text in the first table of the template."""
    if not doc.tables:
        return
    t = doc.tables[0]
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p_style = (p.style.name or "").lower() if p.style else ""
                if title and "title" in p_style and "sub" not in p_style:
                    _replace_text(p, title)
                elif subtitle and "subtitle" in p_style:
                    _replace_text(p, subtitle)


def _replace_text(paragraph, new_text: str) -> None:
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def clear_body_keep_title_block(doc) -> None:
    """Remove all top-level body content except the first table (title block) and sectPr."""
    body = doc.element.body
    seen_title_table = False
    for child in list(body):
        if child.tag == qn("w:tbl"):
            if not seen_title_table:
                seen_title_table = True
                continue
            body.remove(child)
        elif child.tag == qn("w:p"):
            body.remove(child)


def render_table(doc, header_cells: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(header_cells))
    if has_style(doc, "Table Grid"):
        table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(header_cells):
        p = hdr[i].paragraphs[0]
        r = p.add_run(h); r.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, cell in enumerate(row):
            p = table.rows[ri].cells[ci].paragraphs[0]
            add_inline(p, cell)


def parse_and_render(doc, md: str) -> None:
    lines = md.splitlines()
    h1 = style_name(doc, "Heading 1")
    h2 = style_name(doc, "Heading 2")
    h3 = style_name(doc, "Heading 3")
    registry = NumberingRegistry(doc)
    current_list: dict[str, int | None] = {"kind": None, "num_id": None}

    def list_num_id(item_kind: str, level: int) -> int:
        # A single list container keeps one numId for all its levels. The list's
        # "kind" is set by its top-level (level 0) item; nested items reuse that
        # numId so the outer counter is not reset by intervening sub-bullets.
        if level == 0:
            if current_list["kind"] != item_kind or current_list["num_id"] is None:
                current_list["kind"] = item_kind
                current_list["num_id"] = registry.new_list(item_kind)
        elif current_list["num_id"] is None:
            # Nested item with no parent — treat as start of a new list.
            current_list["kind"] = item_kind
            current_list["num_id"] = registry.new_list(item_kind)
        return current_list["num_id"]  # type: ignore[return-value]

    def end_list() -> None:
        current_list["kind"] = None
        current_list["num_id"] = None

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped or stripped == "---":
            i += 1
            continue

        # Markdown table
        if stripped.startswith("|") and stripped.endswith("|"):
            end_list()
            header = [c.strip() for c in stripped.strip("|").split("|")]
            i += 1
            if i < len(lines) and re.match(r"^\s*\|[\s\-:|]+\|\s*$", lines[i]):
                i += 1
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            render_table(doc, header, rows)
            continue

        if stripped.startswith("#### "):
            end_list(); doc.add_paragraph(stripped[5:], style=style_name(doc, "Heading 4")); i += 1; continue
        if stripped.startswith("### "):
            end_list(); doc.add_paragraph(stripped[4:], style=h3); i += 1; continue
        if stripped.startswith("## "):
            end_list(); doc.add_paragraph(stripped[3:], style=h2); i += 1; continue
        if stripped.startswith("# "):
            end_list(); doc.add_paragraph(stripped[2:], style=h1); i += 1; continue

        m = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m:
            indent_spaces = len(m.group(1))
            level = min(indent_spaces // 2, 2)
            p = doc.add_paragraph(style="Normal")
            add_inline(p, m.group(3))
            apply_list_item(p, list_num_id("decimal", level), level)
            i += 1; continue

        m = re.match(r"^(\s*)-\s+(.*)$", line)
        if m:
            indent_spaces = len(m.group(1))
            level = min(indent_spaces // 2, 2)
            p = doc.add_paragraph(style="Normal")
            add_inline(p, m.group(2))
            apply_list_item(p, list_num_id("bullet", level), level)
            i += 1; continue

        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            end_list()
            p = doc.add_paragraph(style="Normal")
            r = p.add_run(stripped.strip("*")); r.italic = True
            i += 1; continue

        end_list()
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.space_after = Pt(BODY_SPACE_AFTER_PT)
        add_inline(p, stripped)
        i += 1


def convert(
    md_path: Path,
    out_path: Path,
    template_path: Path,
    title: str | None = None,
    subtitle: str | None = None,
    strip_front_matter: bool = True,
) -> Path:
    doc = Document(str(template_path))
    make_heading3_bold(doc)
    set_title_block(doc, title, subtitle)
    clear_body_keep_title_block(doc)

    md = md_path.read_text()
    # If the markdown starts with a # H1 + metadata block ending in `---`,
    # strip it so it doesn't duplicate the title block.
    if strip_front_matter and "\n---\n" in md:
        md = md.split("\n---\n", 1)[1]

    parse_and_render(doc, md)
    doc.save(str(out_path))
    return out_path


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("input", type=Path, help="Input markdown file")
    ap.add_argument("output", type=Path, help="Output .docx path")
    ap.add_argument("--title", default=None)
    ap.add_argument("--subtitle", default=None)
    ap.add_argument(
        "--template",
        type=Path,
        default=Path(__file__).with_name("sil-template.docx"),
    )
    ap.add_argument(
        "--keep-front-matter",
        action="store_true",
        help="Do not strip the first '# Title ... ---' block from the markdown.",
    )
    args = ap.parse_args()

    out = convert(
        md_path=args.input,
        out_path=args.output,
        template_path=args.template,
        title=args.title,
        subtitle=args.subtitle,
        strip_front_matter=not args.keep_front_matter,
    )
    print(f"Wrote {out} ({out.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
